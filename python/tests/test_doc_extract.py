"""文档提取测试（M2：TXT/MD/PDF/DOCX，architecture.md §10.4）。"""
from __future__ import annotations

from pathlib import Path

import pytest

from omnisearch.worker.pipeline.doc import DocumentExtractError, extract_text


def test_txt_extract(tmp_path):
    f = tmp_path / "a.txt"
    f.write_text("纯文本内容 hello", encoding="utf-8")
    assert "纯文本内容" in extract_text(str(f))


def test_md_extract(tmp_path):
    f = tmp_path / "b.md"
    f.write_text("# 标题\n\n正文段落", encoding="utf-8")
    text = extract_text(str(f))
    assert "标题" in text and "正文段落" in text


def test_pdf_extract(tmp_path):
    import pymupdf

    pdf = tmp_path / "c.pdf"
    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), "hello pdf content")
    doc.save(str(pdf))
    doc.close()
    assert "hello pdf content" in extract_text(str(pdf))


def test_docx_extract(tmp_path):
    import docx

    f = tmp_path / "d.docx"
    d = docx.Document()
    d.add_paragraph("第一段落")
    d.add_paragraph("第二段落")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "cell1"
    table.rows[0].cells[1].text = "cell2"
    d.save(str(f))
    text = extract_text(str(f))
    assert "第一段落" in text and "第二段落" in text
    assert "cell1 | cell2" in text  # 表格保留原样（不切）


def test_corrupt_pdf_raises(tmp_path):
    f = tmp_path / "bad.pdf"
    f.write_bytes(b"not a real pdf")
    with pytest.raises(DocumentExtractError):
        extract_text(str(f))


def test_unsupported_extension_raises(tmp_path):
    f = tmp_path / "e.xyz"
    f.write_text("x")
    with pytest.raises(DocumentExtractError):
        extract_text(str(f))
