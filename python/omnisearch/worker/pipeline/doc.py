"""文档文本提取（architecture.md §10.4 document pipeline）。

支持 TXT / MD / PDF / DOCX；失败抛 DocumentExtractError → task FAILED（旧数据保留）。
"""
from __future__ import annotations

from pathlib import Path

TEXT_EXTENSIONS = {".txt", ".md", ".markdown"}


class DocumentExtractError(Exception):
    """提取失败（损坏/加密/格式不支持）。"""


def extract_text(path: str) -> str:
    """按扩展名分发提取，返回纯文本。"""
    ext = Path(path).suffix.lower()
    if ext in TEXT_EXTENSIONS:
        return _extract_text_file(path)
    if ext == ".pdf":
        return _extract_pdf(path)
    if ext == ".docx":
        return _extract_docx(path)
    raise DocumentExtractError(f"unsupported document extension: {ext}")


def _extract_text_file(path: str) -> str:
    try:
        return Path(path).read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        raise DocumentExtractError(f"read failed: {exc}") from exc


def _extract_pdf(path: str) -> str:
    try:
        import pymupdf  # PyMuPDF（fitz API 已弃用）

        with pymupdf.open(path) as doc:
            return "\n".join(page.get_text() for page in doc)
    except DocumentExtractError:
        raise
    except Exception as exc:  # noqa: BLE001 —— 损坏/加密/不支持
        raise DocumentExtractError(f"pdf extract failed: {exc}") from exc


def _extract_docx(path: str) -> str:
    try:
        import docx

        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        # 表格保留原样（不切，architecture.md §10.4：表格/代码块不切）
        for table in d.tables:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            if rows:
                parts.append("\n".join(rows))
        return "\n\n".join(parts)
    except DocumentExtractError:
        raise
    except Exception as exc:  # noqa: BLE001
        raise DocumentExtractError(f"docx extract failed: {exc}") from exc
